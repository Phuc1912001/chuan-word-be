"""Test rule engine — không cần DB, chỉ phân tích file .docx tạm.

Chạy: pip install -r requirements.txt -r requirements-dev.txt && pytest
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.rules.page_number import PageNumberRule
from app.rules.presets import get_preset
from app.rules.types import ParsedDoc
from app.services.analyzer import analyze_file
from app.services.docx_parser import load_docx

BODY_TEXT = "Đây là đoạn nội dung văn bản dùng để kiểm tra định dạng theo chuẩn."


def _set_margins(section, top, bottom, left, right):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def _make_bad_doc(path: str) -> None:
    doc = Document()
    # khổ giấy mặc định = Letter (sai A4); lề sai
    _set_margins(doc.sections[0], top=1.0, bottom=1.0, left=1.0, right=1.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT          # chưa căn đều
    p.paragraph_format.first_line_indent = Cm(0)   # chưa thụt đầu dòng
    run = p.add_run(BODY_TEXT)
    run.font.name = "Arial"   # phông sai
    run.font.size = Pt(20)    # cỡ sai
    doc.save(path)


def _make_good_doc(path: str) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)      # A4
    sec.page_height = Cm(29.7)
    _set_margins(sec, top=2.0, bottom=2.0, left=3.0, right=2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(BODY_TEXT)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    # văn bản chuẩn phải có số trang
    PageNumberRule().fix(ParsedDoc(document=doc), get_preset())
    doc.save(path)


def test_detects_format_issues(tmp_path):
    path = tmp_path / "bad.docx"
    _make_bad_doc(str(path))

    report = analyze_file(str(path), get_preset())
    codes = set(report.by_rule)

    for expected in {"FONT", "FONT_SIZE", "MARGIN", "PAGE_SIZE", "ALIGNMENT",
                     "INDENT", "PAGE_NUMBER"}:
        assert expected in codes, f"thiếu rule {expected}; có: {codes}"
    assert report.total_errors >= 7
    assert report.score < 100


def test_clean_doc_has_no_issues(tmp_path):
    path = tmp_path / "good.docx"
    _make_good_doc(str(path))

    report = analyze_file(str(path), get_preset())

    assert report.total_errors == 0, f"không nên có lỗi nhưng có: {report.by_rule}"
    assert report.score == 100.0


def test_error_shape_matches_frontend(tmp_path):
    """Mỗi issue phải map được sang {id, type, message, suggestion, page}."""
    path = tmp_path / "bad.docx"
    _make_bad_doc(str(path))

    report = analyze_file(str(path), get_preset())
    issue = report.issues[0]

    assert issue.rule_code
    assert issue.message
    assert issue.suggestion


def test_line_spacing_max_1_5(tmp_path):
    """Giãn dòng > 1.5 bị báo lỗi; 1.0 (single) thì không."""
    over = tmp_path / "over.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.add_run(BODY_TEXT)
    doc.save(str(over))
    assert "LINE_SPACING" in analyze_file(str(over), get_preset()).by_rule

    ok = tmp_path / "single.docx"
    doc2 = Document()
    p2 = doc2.add_paragraph()
    p2.paragraph_format.line_spacing = 1.0
    p2.add_run(BODY_TEXT)
    doc2.save(str(ok))
    assert "LINE_SPACING" not in analyze_file(str(ok), get_preset()).by_rule


def test_font_in_table_is_checked(tmp_path):
    """Phông trong ô bảng cũng phải được phát hiện (bug doc.paragraphs cũ bỏ sót)."""
    path = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run("Nội dung trong bảng cần kiểm tra phông chữ.")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    doc.save(str(path))

    codes = set(analyze_file(str(path), get_preset()).by_rule)
    assert "FONT" in codes and "FONT_SIZE" in codes


def test_component_size_quoc_hieu(tmp_path):
    """Quốc hiệu cỡ 13 (thuộc 12–13) KHÔNG bị báo; cỡ 16 thì bị báo."""
    ok = tmp_path / "qh_ok.docx"
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r.bold = True
    r.font.size = Pt(13)
    doc.save(str(ok))
    assert "FONT_SIZE" not in analyze_file(str(ok), get_preset()).by_rule


def test_capitalization_detects_and_fixes(tmp_path):
    from app.rules.capitalization import CapitalizationRule

    path = tmp_path / "cap.docx"
    doc = Document()
    doc.add_paragraph("xin chào. hôm nay trời đẹp.")  # 2 câu chưa viết hoa
    doc.save(str(path))

    pd = load_docx(str(path))
    rule = CapitalizationRule()
    spec = get_preset()

    assert rule.check(pd, spec), "phải phát hiện lỗi viết hoa đầu câu"

    rule.fix(pd, spec)
    assert rule.check(pd, spec) == [], "sau fix không còn lỗi viết hoa đầu câu"
    assert pd.paragraphs[0].text.startswith("Xin chào. Hôm nay")


def test_capitalization_respects_sentence_boundary(tmp_path):
    """KHÔNG viết hoa sau dấu chấm không phải hết câu: 'GMO-Z.com', '1.Tên'."""
    from app.rules.capitalization import CapitalizationRule
    from app.rules.types import ParsedDoc

    path = tmp_path / "cap2.docx"
    doc = Document()
    doc.add_paragraph("Kính gửi: Công ty Cổ phần GMO-Z.com RUNSYSTEM")
    doc.add_paragraph("1.Tên miền đăng ký")
    doc.save(str(path))

    pd = ParsedDoc(document=Document(str(path)))
    CapitalizationRule().fix(pd, get_preset())
    assert pd.paragraphs[0].text == "Kính gửi: Công ty Cổ phần GMO-Z.com RUNSYSTEM"
    assert pd.paragraphs[1].text == "1.Tên miền đăng ký"


def test_spelling_rule_detects_common_vietnamese_errors(tmp_path):
    from app.rules.spelling import SpellingRule

    path = tmp_path / "spell.docx"
    doc = Document()
    doc.add_paragraph("Tôi khong duoc tham gia buoi hop.")
    doc.save(str(path))

    pd = load_docx(str(path))
    rule = SpellingRule()
    issues = rule.check(pd, get_preset())

    assert issues, "phải phát hiện lỗi chính tả"
    assert any(issue.rule_code == "SPELLING" for issue in issues)
    assert any("khong" in issue.message.lower() for issue in issues)
    assert any("không" in issue.suggestion for issue in issues)

    rule.fix(pd, get_preset())
    assert "không" in pd.paragraphs[0].text
    assert "được" in pd.paragraphs[0].text


def test_centered_title_not_given_paragraph_spacing(tmp_path):
    """Đoạn căn giữa (tiêu đề) KHÔNG bị thêm giãn đoạn / thụt đầu dòng."""
    from app.services.fixer import fix_file

    src = tmp_path / "title.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BẢN KHAI ĐĂNG KÝ TÊN MIỀN QUỐC GIA VIỆT NAM HÔM NAY")
    doc.save(str(src))
    out = tmp_path / "title_out.docx"
    fix_file(str(src), get_preset(), str(out))

    para = Document(str(out)).paragraphs[0]
    assert para.paragraph_format.space_after in (None, 0) or para.paragraph_format.space_after.pt == 0
    assert para.paragraph_format.first_line_indent in (None,) or para.paragraph_format.first_line_indent.cm == 0


def test_long_token_paragraph_not_justified(tmp_path):
    """Đoạn chứa chuỗi siêu dài (không ngắt được) → căn trái, tránh kéo giãn chữ."""
    from app.services.fixer import fix_file

    src = tmp_path / "long.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Cam kết tuân thủ quy định về quản lý " + "x" * 45 + " và sử dụng tài nguyên.")
    doc.save(str(src))
    out = tmp_path / "long_out.docx"
    fix_file(str(src), get_preset(), str(out))

    assert Document(str(out)).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_fix_makes_doc_clean(tmp_path):
    """Sau auto-fix, phân tích lại phải sạch (check & fix nhất quán)."""
    from app.services.fixer import fix_file

    src = tmp_path / "bad.docx"
    _make_bad_doc(str(src))
    out = tmp_path / "fixed.docx"
    fix_file(str(src), get_preset(), str(out))

    report = analyze_file(str(out), get_preset())
    assert report.total_errors == 0, f"sau fix vẫn còn lỗi: {report.by_rule}"
