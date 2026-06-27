"""Rule số trang (PAGE_NUMBER) — văn bản phải được đánh số trang.

CHECK: nếu KHÔNG footer nào của tài liệu chứa trường PAGE → báo thiếu số trang.
FIX: chèn số trang căn giữa vào footer (Times New Roman, cỡ trong khoảng chuẩn).
"""

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.rules.docx_iter import set_run_font_name
from app.rules.types import COMP_NOI_DUNG, Issue, ParsedDoc, TemplateSpec


def _has_page_field(part) -> bool:
    el = part._element
    for fld in el.iter(qn("w:fldSimple")):
        if "PAGE" in (fld.get(qn("w:instr")) or "").upper():
            return True
    for instr in el.iter(qn("w:instrText")):
        if instr.text and "PAGE" in instr.text.upper():
            return True
    return False


def _unique_footers(document):
    seen: set[int] = set()
    out = []
    for sec in document.sections:
        f = sec.footer
        key = id(f._element)
        if key not in seen:
            seen.add(key)
            out.append(f)
    return out


def _add_page_number(paragraph, font_name: str, size_pt: float) -> None:
    from docx.shared import Pt

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font_name(run, font_name)
    run.font.size = Pt(size_pt)


class PageNumberRule:
    code = "PAGE_NUMBER"
    name = "Số trang"

    def check(self, doc: ParsedDoc, spec: TemplateSpec) -> list[Issue]:
        document = doc.document
        if any(_has_page_field(f) for f in _unique_footers(document)):
            return []
        return [Issue(
            rule_code=self.code,
            message="Văn bản chưa được đánh số trang",
            suggestion="Đánh số trang ở chân trang, căn giữa",
        )]

    def fix(self, doc: ParsedDoc, spec: TemplateSpec) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = doc.document
        if any(_has_page_field(f) for f in _unique_footers(document)):
            return

        sec = document.sections[0]
        footer = sec.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size_pt = spec.component(COMP_NOI_DUNG).fix_size_pt
        _add_page_number(para, spec.font.name, size_pt)
