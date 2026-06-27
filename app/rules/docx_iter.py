"""Tiện ích duyệt & chỉnh OOXML dùng chung cho rule engine.

LÝ DO TỒN TẠI: `Document.paragraphs` của python-docx CHỈ trả đoạn ở thân tài
liệu — bỏ qua đoạn trong bảng (table), header và footer. Mọi rule trước đây
duyệt `doc.paragraphs` nên KHÔNG bao giờ chạm tới nội dung trong bảng/đầu/chân
trang. `iter_all_paragraphs` đi đệ quy qua tất cả các vùng đó.
"""

from __future__ import annotations

from typing import Iterator

from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.rules.types import ParsedDoc


def _document(doc):
    return doc.document if isinstance(doc, ParsedDoc) else doc


def _iter_block_items(parent) -> Iterator[object]:
    """Yield Paragraph và Table con TRỰC TIẾP của `parent`, đúng thứ tự tài liệu.

    `parent` là Document, _Cell, _Header hoặc _Footer (đều có `._element`/`.element`).
    """
    elm = getattr(parent, "element", None)
    if elm is not None and elm.tag == qn("w:document"):
        parent_elm = elm.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _walk(parent) -> Iterator[Paragraph]:
    for block in _iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block
        else:  # Table → đệ quy vào từng ô (xử lý cả bảng lồng bảng)
            for row in block.rows:
                for cell in row.cells:
                    yield from _walk(cell)


def _walk_with_flag(parent, in_table: bool) -> Iterator[tuple]:
    """Như _walk nhưng kèm cờ in_table (đoạn có nằm trong bảng không)."""
    for block in _iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block, in_table
        else:
            for row in block.rows:
                for cell in row.cells:
                    yield from _walk_with_flag(cell, True)


def iter_top_level_paragraphs(doc) -> list[Paragraph]:
    """CHỈ đoạn ở cấp thân tài liệu, KHÔNG vào trong bảng — dùng cho rule bố cục
    (thụt dòng, giãn đoạn) vốn không được áp lên ô bảng."""
    return list(_document(doc).paragraphs)


def iter_body_with_flag(doc) -> Iterator[tuple]:
    """Yield (paragraph, in_table) cho mọi đoạn thân, đúng thứ tự tài liệu."""
    yield from _walk_with_flag(_document(doc), False)


def _walk_with_flag(parent, in_table: bool) -> Iterator[tuple]:
    """Như _walk nhưng kèm cờ in_table (đoạn có nằm trong bảng không)."""
    for block in _iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block, in_table
        else:
            for row in block.rows:
                for cell in row.cells:
                    yield from _walk_with_flag(cell, True)


def _header_footer_parts(section) -> Iterator[object]:
    for attr in (
        "header",
        "footer",
        "first_page_header",
        "first_page_footer",
        "even_page_header",
        "even_page_footer",
    ):
        part = getattr(section, attr, None)
        if part is not None:
            yield part


def iter_header_footer_paragraphs(doc) -> Iterator[Paragraph]:
    """Đoạn trong header/footer của mọi section (đã khử trùng lặp khi link)."""
    document = _document(doc)
    seen: set[int] = set()
    for section in document.sections:
        for part in _header_footer_parts(section):
            key = id(part._element)
            if key in seen:
                continue
            seen.add(key)
            yield from _walk(part)


def iter_all_paragraphs(doc, include_headers_footers: bool = True) -> Iterator[Paragraph]:
    """Mọi đoạn văn của tài liệu: thân + bảng (lồng nhau) + header/footer.

    THỨ TỰ ỔN ĐỊNH: luôn là [tất cả đoạn thân] rồi tới [header/footer]. Nhờ vậy
    classifier có thể map nhãn theo VỊ TRÍ (id của proxy lxml không bền vững)."""
    yield from _walk(_document(doc))
    if include_headers_footers:
        yield from iter_header_footer_paragraphs(doc)


def iter_body_paragraphs(doc) -> list[Paragraph]:
    """Chỉ đoạn ở thân tài liệu (kể cả trong bảng), theo thứ tự — dùng cho classifier."""
    return list(_walk(_document(doc)))


# --- Chỉnh font ở mức run (đặt đủ 4 thuộc tính rFonts) ---

def set_run_font_name(run, name: str) -> None:
    """Đặt phông cho run đủ ascii/hAnsi/cs/eastAsia.

    `run.font.name = x` của python-docx chỉ set ascii+hAnsi → một số file
    (có w:cs/w:eastAsia trỏ phông khác) vẫn render sai. Hàm này set cả 4.
    """
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def run_font_name(run) -> str | None:
    """Tên phông tường minh của run (đọc ascii rồi tới các biến thể)."""
    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.rFonts
    if rfonts is None:
        return None
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        val = rfonts.get(qn(attr))
        if val:
            return val
    return None


def effective_size_pt(run, para, document) -> float | None:
    """Cỡ chữ HIỆU LỰC của run: explicit → style của đoạn (theo chuỗi base) → Normal."""
    if run.font.size is not None:
        return run.font.size.pt
    style = getattr(para, "style", None)
    while style is not None:
        size = style.font.size
        if size is not None:
            return size.pt
        style = getattr(style, "base_style", None)
    try:
        nz = document.styles["Normal"].font.size
        if nz is not None:
            return nz.pt
    except KeyError:
        pass
    return None
